import os
import uuid
import json
import re
import boto3
from botocore.exceptions import BotoCoreError, ClientError
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from models.models import Resume, Candidate, Job
from dotenv import load_dotenv

load_dotenv()

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}

S3_BUCKET        = os.getenv("S3_BUCKET")
S3_REGION        = os.getenv("S3_REGION", "ap-south-1")
S3_ACCESS_KEY_ID = os.getenv("S3_ACCESS_KEY_ID")
S3_SECRET_ACCESS = os.getenv("S3_SECRET_ACCESS")
BEDROCK_MODEL    = "meta.llama3-8b-instruct-v1:0"

SKILLS_POOL = [
    "python", "javascript", "typescript", "react", "node.js", "nodejs", "sql",
    "postgresql", "mysql", "mongodb", "aws", "docker", "kubernetes", "machine learning",
    "tensorflow", "pytorch", "fastapi", "django", "flask", "spring boot", "java",
    "go", "golang", "graphql", "redis", "kafka", "spark", "ci/cd", "terraform",
    "figma", "product management", "data analysis", "c++", "c#", "rust", "php",
    "vue", "angular", "linux", "git", "agile", "scrum",
]


# ── AWS clients ────────────────────────────────────────────────────────────

def _s3_client():
    return boto3.client(
        "s3",
        region_name=S3_REGION,
        aws_access_key_id=S3_ACCESS_KEY_ID,
        aws_secret_access_key=S3_SECRET_ACCESS,
    )


def _bedrock_client():
    return boto3.client(
        "bedrock-runtime",
        region_name=S3_REGION,
        aws_access_key_id=S3_ACCESS_KEY_ID,
        aws_secret_access_key=S3_SECRET_ACCESS,
    )


# ── S3 helpers ─────────────────────────────────────────────────────────────

def _upload_to_s3(file_bytes: bytes, s3_key: str, content_type: str) -> str:
    if not S3_BUCKET:
        raise HTTPException(status_code=500, detail="S3_BUCKET not configured.")
    try:
        client = _s3_client()
        client.put_object(Bucket=S3_BUCKET, Key=s3_key, Body=file_bytes, ContentType=content_type)
        return client.generate_presigned_url(
            "get_object",
            Params={"Bucket": S3_BUCKET, "Key": s3_key},
            ExpiresIn=365 * 24 * 3600,
        )
    except (BotoCoreError, ClientError) as e:
        raise HTTPException(status_code=500, detail=f"S3 upload failed: {str(e)}")


# ── Text extraction ────────────────────────────────────────────────────────

def _extract_text_from_bytes(file_bytes: bytes, ext: str) -> str:
    if ext == ".pdf":
        # Try pdfplumber first (most reliable)
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = " ".join(page.extract_text() or "" for page in pdf.pages)
            if text.strip():
                return text
        except Exception:
            pass
        # Fallback to PyPDF2
        try:
            import PyPDF2, io
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = " ".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text
        except Exception:
            pass
        return ""
    elif ext in (".docx", ".doc"):
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # Extract from paragraphs
            para_text = " ".join(p.text for p in doc.paragraphs)
            
            # Extract from tables (many resumes use tables for layout)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
            
            # Combine both sources
            full_text = para_text + " " + " ".join(table_text)
            return full_text.strip()
        except Exception:
            return ""
    return ""


def _parse_skills(text: str) -> list[str]:
    text_lower = text.lower()
    return [skill for skill in SKILLS_POOL if skill in text_lower]


def _default_resume_profile() -> dict:
    return {
        "full_name": "",
        "headline": "",
        "summary": "",
        "contact": {
            "email": "",
            "phone": "",
            "address": "",
            "location": "",
        },
        "links": {
            "linkedin": "",
            "github": "",
            "portfolio": "",
            "website": "",
        },
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": [],
        "achievements": [],
        "languages": [],
        "raw_sections": {
            "summary": "",
            "experience": "",
            "projects": "",
            "education": "",
            "skills": "",
            "certifications": "",
            "other": "",
        },
    }


def _merge_resume_profile(base: dict, data: dict) -> dict:
    merged = base.copy()
    for key in ["full_name", "headline", "summary"]:
        value = data.get(key)
        if isinstance(value, str) and value.strip():
            merged[key] = value.strip()

    for section in ["contact", "links", "raw_sections"]:
        section_data = base.get(section, {}).copy()
        incoming = data.get(section, {}) or {}
        if isinstance(incoming, dict):
            for key, value in incoming.items():
                if isinstance(value, str):
                    section_data[key] = value.strip()
        merged[section] = section_data

    for list_key in ["skills", "experience", "projects", "education", "certifications", "achievements", "languages"]:
        value = data.get(list_key)
        if isinstance(value, list) and value:
            merged[list_key] = value

    return merged


def _heuristic_resume_profile(text: str) -> dict:
    profile = _default_resume_profile()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # Extract contact information
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}", text)
    
    # Improved URL extraction (without https://)
    linkedin_pattern = r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s)]+"
    github_pattern = r"(?:https?://)?(?:www\.)?github\.com/[^\s)]+"
    website_pattern = r"(?:https?://)?(?:www\.)?[\w\-\.]+\.[\w]{2,}(?:/[^\s)]*)?"
    
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    
    # Find website that's not linkedin or github
    websites = re.findall(website_pattern, text, re.IGNORECASE)
    portfolio = None
    for site in websites:
        site_lower = site.lower()
        if 'linkedin' not in site_lower and 'github' not in site_lower and 'gmail' not in site_lower:
            portfolio = site
            break

    profile["contact"]["email"] = email_match.group(0) if email_match else ""
    profile["contact"]["phone"] = phone_match.group(0) if phone_match else ""
    profile["links"]["linkedin"] = linkedin_match.group(0) if linkedin_match else ""
    profile["links"]["github"] = github_match.group(0) if github_match else ""
    profile["links"]["portfolio"] = portfolio or ""
    profile["links"]["website"] = portfolio or ""

    # Extract candidate name
    candidate_name = ""
    for line in lines[:8]:
        lower_line = line.lower()
        if any(token in lower_line for token in ["@", "linkedin.com", "github.com", "experience", "education", "skills", "summary"]):
            continue
        if 2 <= len(line.split()) <= 5 and sum(char.isalpha() for char in line) >= 4:
            candidate_name = line
            break
    profile["full_name"] = candidate_name

    # Enhanced section headings
    section_headings = {
        "summary": ["summary", "professional summary", "profile", "about", "objective"],
        "experience": ["experience", "work experience", "employment history", "professional experience", "work history"],
        "projects": ["projects", "project experience", "selected projects", "key projects", "major projects"],
        "education": ["education", "academics", "academic background", "qualifications"],
        "skills": ["skills", "technical skills", "core skills", "competencies", "technologies"],
        "certifications": ["certifications", "certificates", "licenses"],
        "achievements": ["achievements", "accomplishments", "awards", "honors", "recognition"],
    }

    # Parse sections
    active_section = None
    section_lines = {key: [] for key in section_headings}
    other_lines = []

    for line in lines:
        normalized = re.sub(r"[:\-]+$", "", line.strip().lower())
        matched_section = next((section for section, headings in section_headings.items() if normalized in headings), None)
        if matched_section:
            active_section = matched_section
            continue

        if active_section and active_section in section_lines:
            section_lines[active_section].append(line)
        else:
            other_lines.append(line)

    profile["summary"] = " ".join(other_lines[:3]).strip()
    profile["raw_sections"] = {
        "summary": " ".join(section_lines["summary"]).strip(),
        "experience": " ".join(section_lines["experience"]).strip(),
        "projects": " ".join(section_lines["projects"]).strip(),
        "education": " ".join(section_lines["education"]).strip(),
        "skills": " ".join(section_lines["skills"]).strip(),
        "certifications": " ".join(section_lines["certifications"]).strip(),
        "achievements": " ".join(section_lines["achievements"]).strip(),
        "other": " ".join(other_lines).strip(),
    }
    
    # Parse skills
    profile["skills"] = _parse_skills(text)
    
    # Parse experience with highlights
    profile["experience"] = _parse_experience_section(section_lines["experience"])
    
    # Calculate total years of experience from work history
    total_experience_years = _calculate_experience_years(profile["experience"])
    profile["total_experience_years"] = total_experience_years
    
    # Parse projects with details
    profile["projects"] = _parse_projects_section(section_lines["projects"])
    
    # Parse education
    profile["education"] = _parse_education_section(section_lines["education"])
    
    # Parse certifications (split by line or bullet)
    certs_text = profile["raw_sections"]["certifications"]
    if certs_text:
        # Split by common separators
        cert_lines = re.split(r'[•\n\|\-]', certs_text)
        profile["certifications"] = [cert.strip() for cert in cert_lines if cert.strip() and len(cert.strip()) > 5]
    
    # Parse achievements
    achievements_text = profile["raw_sections"]["achievements"]
    if achievements_text:
        # Split by bullet points or lines
        achievement_lines = re.split(r'[•\n]', achievements_text)
        profile["achievements"] = [ach.strip() for ach in achievement_lines if ach.strip() and len(ach.strip()) > 10]
    
    return profile


def _calculate_experience_years(experience_list: list) -> float:
    """Calculate total years of experience from work history"""
    if not experience_list:
        return 0.0
    
    total_months = 0
    current_year = 2026  # From system context
    current_month = 6     # June 2026
    
    for exp in experience_list:
        if not isinstance(exp, dict):
            continue
        
        # Look for date patterns in the experience data
        exp_str = f"{exp.get('title', '')} {exp.get('company', '')} {exp.get('details', '')}".lower()
        
        # Common date patterns: "Jun 2025 - Aug 2025", "Oct 2023 - Apr 2024", "2020-2023"
        date_patterns = [
            r'(\w{3})\s+(\d{4})\s*[-–]\s*(\w{3})\s+(\d{4})',  # Jun 2025 - Aug 2025
            r'(\w{3})\s+(\d{4})\s*[-–]\s*present',             # Jun 2025 - Present
            r'(\d{4})\s*[-–]\s*(\d{4})',                       # 2020-2023
            r'(\d{4})\s*[-–]\s*present',                       # 2020-Present
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, exp_str, re.IGNORECASE)
            if matches:
                match = matches[0]
                
                try:
                    if len(match) == 4:  # Month-Year format (Jun 2025 - Aug 2025)
                        start_month_str, start_year, end_month_str, end_year = match
                        month_map = {
                            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
                        }
                        start_month = month_map.get(start_month_str[:3].lower(), 1)
                        end_month = month_map.get(end_month_str[:3].lower(), 12)
                        start_year = int(start_year)
                        end_year = int(end_year)
                        
                        # Calculate months
                        months = (end_year - start_year) * 12 + (end_month - start_month) + 1
                        total_months += max(0, months)
                        
                    elif len(match) == 2 and 'present' not in pattern:  # Year format (2020-2023)
                        start_year, end_year = int(match[0]), int(match[1])
                        months = (end_year - start_year) * 12
                        total_months += max(0, months)
                        
                    elif 'present' in pattern:  # Currently working
                        if len(match) == 2:  # Jun 2025 - Present
                            start_month_str, start_year = match
                            month_map = {
                                'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                                'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
                            }
                            start_month = month_map.get(start_month_str[:3].lower(), 1)
                            start_year = int(start_year)
                            months = (current_year - start_year) * 12 + (current_month - start_month) + 1
                            total_months += max(0, months)
                        else:  # 2020 - Present
                            start_year = int(match[0])
                            months = (current_year - start_year) * 12
                            total_months += max(0, months)
                            
                except (ValueError, KeyError):
                    continue
                
                break  # Found a date pattern, move to next experience
    
    # Convert months to years (rounded to 1 decimal)
    years = round(total_months / 12, 1)
    return years


def _parse_experience_section(lines: list) -> list:
    """Parse experience section into structured format"""
    experiences = []
    if not lines:
        return experiences
    
    current_exp = None
    
    for line in lines:
        # Check if line looks like a job title (capitalized, not too long, not a bullet)
        if line and not line.startswith(('•', '-', '✓')) and len(line) < 100:
            # Likely a new job entry
            if current_exp:
                experiences.append(current_exp)
            
            # Try to extract company and title
            parts = re.split(r'\s+at\s+|\s+@\s+', line, flags=re.IGNORECASE)
            if len(parts) == 2:
                title, company = parts
            else:
                title = line
                company = ""
            
            current_exp = {
                "title": title.strip(),
                "company": company.strip(),
                "highlights": []
            }
        elif current_exp and line:
            # Add as highlight
            clean_line = line.lstrip('•-✓ ').strip()
            if clean_line:
                current_exp["highlights"].append(clean_line)
    
    if current_exp:
        experiences.append(current_exp)
    
    return experiences if experiences else [{"details": " ".join(lines)}]


def _parse_projects_section(lines: list) -> list:
    """Parse projects section into structured format"""
    projects = []
    if not lines:
        return projects
    
    current_project = None
    
    for line in lines:
        # Check if line looks like a project name (short, not a bullet)
        if line and not line.startswith(('•', '-', '✓', 'Tech:')) and len(line) < 80 and len(line.split()) <= 6:
            # Likely a new project
            if current_project:
                projects.append(current_project)
            
            current_project = {
                "name": line.strip(),
                "description": "",
                "tech_stack": []
            }
        elif current_project:
            # Check if it's a tech stack line
            if line.lower().startswith('tech:') or line.lower().startswith('technologies:'):
                tech_text = re.sub(r'^(tech:|technologies:)', '', line, flags=re.IGNORECASE).strip()
                current_project["tech_stack"] = [t.strip() for t in re.split(r'[,;|]', tech_text) if t.strip()]
            else:
                # Add to description
                clean_line = line.lstrip('•-✓ ').strip()
                if clean_line:
                    if current_project["description"]:
                        current_project["description"] += " " + clean_line
                    else:
                        current_project["description"] = clean_line
    
    if current_project:
        projects.append(current_project)
    
    return projects if projects else [{"details": " ".join(lines)}]


def _parse_education_section(lines: list) -> list:
    """Parse education section into structured format"""
    education = []
    if not lines:
        return education
    
    for line in lines:
        if line and len(line) > 10:
            # Try to extract degree and institution
            parts = re.split(r'\s+-\s+|\s+from\s+|\s+at\s+', line, flags=re.IGNORECASE)
            if len(parts) >= 2:
                education.append({
                    "degree": parts[0].strip(),
                    "institution": parts[1].strip() if len(parts) > 1 else ""
                })
            else:
                education.append({"degree": line.strip(), "institution": ""})
    
    return education if education else [{"details": " ".join(lines)}]


def _run_resume_profile_extraction(resume_text: str) -> dict:
    prompt = f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>
You are a resume parsing engine.
Return ONLY a valid JSON object and no extra text.
Extract structured resume data from the provided text.
<|eot_id|><|start_header_id|>user<|end_header_id|>

Parse this resume text into JSON.

Resume:
{resume_text[:5000]}

Return exactly this JSON structure:
{{
  "full_name": "",
  "headline": "",
  "summary": "",
  "contact": {{
    "email": "",
    "phone": "",
    "address": "",
    "location": ""
  }},
  "links": {{
    "linkedin": "",
    "github": "",
    "portfolio": "",
    "website": ""
  }},
  "skills": [""],
  "experience": [{{
    "title": "",
    "company": "",
    "location": "",
    "start_date": "",
    "end_date": "",
    "highlights": [""]
  }}],
  "projects": [{{
    "name": "",
    "description": "",
    "tech_stack": [""],
    "link": ""
  }}],
  "education": [{{
    "institution": "",
    "degree": "",
    "field": "",
    "start_date": "",
    "end_date": "",
    "details": [""]
  }}],
  "certifications": [""],
  "achievements": [""],
  "languages": [""],
  "raw_sections": {{
    "summary": "",
    "experience": "",
    "projects": "",
    "education": "",
    "skills": "",
    "certifications": "",
    "other": ""
  }}
}}
<|eot_id|><|start_header_id|>assistant<end_header_id|>
"""

    try:
        client = _bedrock_client()
        response = client.invoke_model(
            modelId=BEDROCK_MODEL,
            contentType="application/json",
            accept="application/json",
            body=json.dumps({
                "prompt": prompt,
                "max_gen_len": 2048,
                "temperature": 0.1,
                "top_p": 0.9,
            }),
        )
        raw = json.loads(response["body"].read())
        generation = raw.get("generation", "")
        start = generation.find("{")
        end = generation.rfind("}") + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON found in model response")
        return _merge_resume_profile(_default_resume_profile(), json.loads(generation[start:end]))
    except (BotoCoreError, ClientError, json.JSONDecodeError, ValueError):
        return _heuristic_resume_profile(resume_text)


# ── Bedrock ATS check ──────────────────────────────────────────────────────

def _run_ats_check(resume_text: str, job_description: str = "") -> dict:
    jd_section = (
        f"\n\nJob Description to match against:\n{job_description.strip()}"
        if job_description.strip()
        else ""
    )

    prompt = f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>
You are an expert ATS (Applicant Tracking System) evaluator and career coach.
Analyze the resume and return ONLY a valid JSON object — no extra text, no markdown, no explanation outside the JSON.
<|eot_id|><|start_header_id|>user<|end_header_id|>

Analyze this resume for ATS compatibility{' and match against the provided job description' if job_description.strip() else ''}.

Resume:
{resume_text[:3000]}
{jd_section}

Return exactly this JSON structure:
{{
  "ats_score": <integer 0-100>,
  "breakdown": {{
    "keyword_match": <integer 0-100>,
    "formatting": <integer 0-100>,
    "experience_relevance": <integer 0-100>,
    "skills_coverage": <integer 0-100>,
    "education_fit": <integer 0-100>
  }},
  "strengths": [<3-5 short strings>],
  "improvements": [<4-6 short actionable strings>],
  "missing_keywords": [<up to 8 important missing keywords>],
  "summary": "<2 sentence overall assessment>"
}}
<|eot_id|><|start_header_id|>assistant<|end_header_id|>
"""

    try:
        client = _bedrock_client()
        response = client.invoke_model(
            modelId=BEDROCK_MODEL,
            contentType="application/json",
            accept="application/json",
            body=json.dumps({
                "prompt": prompt,
                "max_gen_len": 1024,
                "temperature": 0.1,
                "top_p": 0.9,
            }),
        )
        raw = json.loads(response["body"].read())
        generation = raw.get("generation", "")

        # Extract JSON from the generation
        start = generation.find("{")
        end   = generation.rfind("}") + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON found in model response")

        return json.loads(generation[start:end])

    except (BotoCoreError, ClientError) as e:
        raise HTTPException(status_code=500, detail=f"Bedrock error: {str(e)}")
    except (json.JSONDecodeError, ValueError):
        # Return a safe fallback if model output is malformed
        return {
            "ats_score": 0,
            "breakdown": {
                "keyword_match": 0, "formatting": 0, "experience_relevance": 0,
                "skills_coverage": 0, "education_fit": 0,
            },
            "strengths": [],
            "improvements": ["Could not parse AI response. Please try again."],
            "missing_keywords": [],
            "summary": "Analysis failed. Please re-upload your resume.",
        }


# ── Public functions ───────────────────────────────────────────────────────

def ats_check(file: UploadFile, job_description: str = "") -> dict:
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are accepted.")

    # Ensure we read from the beginning
    try:
        file.file.seek(0)
    except Exception:
        pass  # Some file-like objects don't support seek
    
    file_bytes = file.file.read()
    
    # Reset pointer again after reading
    try:
        file.file.seek(0)
    except Exception:
        pass

    if not file_bytes:
        raise HTTPException(status_code=400, detail="File is empty or could not be read.")

    raw_text = _extract_text_from_bytes(file_bytes, ext)

    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from the file. Ensure it is not a scanned image.")

    result = _run_ats_check(raw_text, job_description)
    result["filename"]       = file.filename
    result["extracted_text"] = raw_text[:500]  # Limit to 500 chars to keep response smaller
    result["skills_found"]   = _parse_skills(raw_text)
    return result


def upload_resume(db: Session, candidate_id: int, file: UploadFile) -> dict:
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are accepted.")

    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found.")

    # Ensure we read from the beginning
    try:
        file.file.seek(0)
    except Exception:
        pass
    
    file_bytes = file.file.read()
    
    if not file_bytes:
        raise HTTPException(status_code=400, detail="File is empty or could not be read.")
    
    content_type = "application/pdf" if ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    s3_key       = f"resumes/candidate_{candidate_id}/{uuid.uuid4().hex}{ext}"
    s3_url       = _upload_to_s3(file_bytes, s3_key, content_type)

    raw_text     = _extract_text_from_bytes(file_bytes, ext)
    parsed_skills = _parse_skills(raw_text)
    parsed_profile = _run_resume_profile_extraction(raw_text)

    resume = Resume(
        candidate_id=candidate_id,
        filename=file.filename,
        file_path=s3_url,
        parsed_skills=parsed_skills,
        parsed_profile=parsed_profile,
        raw_text=raw_text[:5000],
    )
    db.add(resume)

    candidate.resume_path = s3_url
    if parsed_skills:
        candidate.skills = parsed_skills
    
    # Update experience from calculated years
    if parsed_profile and parsed_profile.get('total_experience_years'):
        candidate.experience = int(parsed_profile['total_experience_years'])

    db.commit()
    db.refresh(resume)
    return {
        "id": resume.id,
        "filename": resume.filename,
        "file_path": resume.file_path,
        "parsed_skills": resume.parsed_skills,
        "parsed_profile": resume.parsed_profile,
        "raw_text": resume.raw_text,
        "uploaded_at": str(resume.uploaded_at),
    }


def match_skills(db: Session, candidate_id: int, job_id: int) -> dict:
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found.")

    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")

    candidate_skills = set(s.lower() for s in (candidate.skills or []))
    job_skills       = set(s.lower() for s in (job.skills or []))
    matched          = sorted(candidate_skills & job_skills)
    missing          = sorted(job_skills - candidate_skills)
    score            = round((len(matched) / len(job_skills) * 100) if job_skills else 0, 1)

    return {
        "candidate_id": candidate_id,
        "job_id": job_id,
        "match_score": score,
        "matched_skills": matched,
        "missing_skills": missing,
        "total_job_skills": len(job_skills),
    }


def get_resumes(db: Session, candidate_id: int):
    return db.query(Resume).filter(Resume.candidate_id == candidate_id).all()
