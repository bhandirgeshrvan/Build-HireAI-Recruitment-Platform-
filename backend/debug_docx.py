"""Debug script to test DOCX reading"""
import docx
import io

file_path = "/home/ideabliss/Hiring-platform/Build-HireAI-Recruitment-Platform-/resume/Aditya_K_Shinde_Resume_ATS.docx"

print(f"Testing file: {file_path}")

# Read file
with open(file_path, 'rb') as f:
    file_bytes = f.read()

print(f"File size: {len(file_bytes)} bytes")

# Try to extract text
try:
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # Extract from paragraphs
    para_text = " ".join(p.text for p in doc.paragraphs)
    print(f"Paragraph text length: {len(para_text)} characters")
    
    # Extract from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    table_text = " ".join(table_text)
    print(f"Table text length: {len(table_text)} characters")
    
    # Combine
    full_text = para_text + " " + table_text
    print(f"\nTotal text length: {len(full_text.strip())} characters")
    print(f"\nFirst 500 characters:\n{full_text[:500]}")
    
    if full_text.strip():
        print(f"\n✓ Success! Text extraction works.")
    else:
        print(f"\n✗ Warning: No text extracted")
except Exception as e:
    print(f"✗ Error: {e}")
    import traceback
    traceback.print_exc()
